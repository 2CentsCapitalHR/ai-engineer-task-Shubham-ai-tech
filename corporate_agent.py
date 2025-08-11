import os
import json
from docx import Document
from langchain_google_genai import GoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.prompts import PromptTemplate
from langchain.schema.runnable import RunnablePassthrough
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from typing import List

# --- PYDANTIC MODELS FOR OUTPUT PARSING ---
class Issue(BaseModel):
    issue: str = Field(description="Description of the compliance issue.")
    severity: str = Field(description="High, Medium, or Low")
    suggestion: str = Field(description="How to fix the issue.")

class ValidationResult(BaseModel):
    issues: List[Issue] = Field(description="A list of compliance issues found.")


# --- CONFIGURATION ---
DB_FAISS_PATH = 'vectorstore'
OUTPUT_PATH = 'outputs'
os.makedirs(OUTPUT_PATH, exist_ok=True)

# Define the checklist for company incorporation
PROCESS_CHECKLISTS = {
    "Company Incorporation": {
        "required_docs": [
            "Articles of Association", "Memorandum of Association", "Incorporation Application Form",
            "UBO Declaration Form", "Register of Members and Directors"
        ],
        "doc_keywords": {
            "Articles of Association": ["articles of association"],
            "Memorandum of Association": ["memorandum of association"],
            "Incorporation Application Form": ["application for incorporation"],
            "UBO Declaration Form": ["beneficial owner", "ubo declaration"],
            "Register of Members and Directors": ["register of members", "register of directors"]
        }
    }
}

# --- LOAD RAG AND LLM ---
embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2', model_kwargs={'device': 'cpu'})
db = FAISS.load_local(DB_FAISS_PATH, embeddings, allow_dangerous_deserialization=True)
retriever = db.as_retriever(search_kwargs={'k': 5}) # Increased k for more context
load_dotenv()

google_api_key = os.getenv("GOOGLE_API_KEY")
client_options_dict = {"api_key": google_api_key}

llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash",
    client_options=client_options_dict
)

# --- PROMPT TEMPLATE ---
parser = PydanticOutputParser(pydantic_object=ValidationResult)

prompt_template = """
You are an expert ADGM compliance officer. Analyze the following document text, which may contain multiple clauses or paragraphs, based ONLY on the provided ADGM regulations context.

{format_instructions}

CONTEXT:
{context}

DOCUMENT TEXT:
{question}

Analyze the document text and provide your response as a single JSON object listing all identified issues from the entire text.
"""
prompt = PromptTemplate(
    template=prompt_template,
    input_variables=['context', 'question'],
    partial_variables={"format_instructions": parser.get_format_instructions()}
)

# --- RAG CHAIN ---
rag_chain = (
    {"context": retriever, "question": RunnablePassthrough()}
    | prompt
    | llm
    | parser
)

# --- CORE FUNCTIONS ---
def identify_doc_type(text):
    text = text.lower()
    for doc_type, keywords in PROCESS_CHECKLISTS["Company Incorporation"]["doc_keywords"].items():
        if any(keyword in text for keyword in keywords):
            return doc_type
    return "Unknown Document"

def analyze_documents(uploaded_file_paths):
    # 1. Document Identification and Checklist Verification
    doc_texts = {}
    identified_docs = set()
    for file_path in uploaded_file_paths:
        doc = Document(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        doc_type = identify_doc_type(full_text[:1000])
        doc_texts[file_path] = doc
        if doc_type != "Unknown Document":
            identified_docs.add(doc_type)

    checklist = PROCESS_CHECKLISTS["Company Incorporation"]
    required = set(checklist["required_docs"])
    missing_docs = list(required - identified_docs)

    # 2. Red Flag Analysis and Commenting
    all_issues = []
    output_docx_path = None
    
    for file_path, doc in doc_texts.items():
        issues_found_in_doc = False
        
        # --- BATCHING LOGIC (OPTION 1) IMPLEMENTED HERE ---
        
        # Define the size of each chunk (number of paragraphs)
        CHUNK_SIZE = 7
        
        # First, gather all paragraphs that are long enough to be analyzed
        paragraphs_to_process = [p for p in doc.paragraphs if len(p.text.strip()) > 50]
        
        # Loop through the collected paragraphs in chunks
        for i in range(0, len(paragraphs_to_process), CHUNK_SIZE):
            
            # Get the slice of paragraph objects for the current chunk
            current_chunk_of_paragraphs = paragraphs_to_process[i:i + CHUNK_SIZE]
            
            # Combine their text to create a single context for the LLM
            chunk_text = "\n---\n".join([p.text for p in current_chunk_of_paragraphs])
            
            # The first paragraph of the chunk will be our anchor for adding comments
            anchor_paragraph = current_chunk_of_paragraphs[0]
            
            try:
                # Invoke the RAG chain only ONCE for the entire chunk
                result_object = rag_chain.invoke(chunk_text)
                
                if result_object.issues:
                    consolidated_comment = ""
                    for issue in result_object.issues:
                        consolidated_comment += f"Issue: {issue.suggestion} (Severity: {issue.severity})\n"
                        
                        # Add each individual issue to the main JSON report
                        all_issues.append({
                            "document": os.path.basename(file_path),
                            "section": f"Analysis of chunk starting with: '{anchor_paragraph.text[:60]}...'",
                            "issue": issue.issue,
                            "severity": issue.severity,
                            "suggestion": issue.suggestion
                        })
                    
                    # Add the consolidated comment as a single, bolded run to the anchor paragraph
                    if consolidated_comment:
                        anchor_paragraph.add_run(f"\n[AI REVIEW FOR THIS SECTION]:\n{consolidated_comment}").bold = True
                        issues_found_in_doc = True

            except Exception as e:
                print(f"Skipping chunk due to processing error: {e}")
                continue
        
        # --- END OF BATCHING LOGIC ---
        
        if issues_found_in_doc:
            filename = os.path.basename(file_path)
            output_docx_path = os.path.join(OUTPUT_PATH, f"reviewed_{filename}")
            doc.save(output_docx_path)

    # 3. Generate Final Report
    summary_report = {
        "process": "Company Incorporation",
        "documents_uploaded": len(identified_docs),
        "required_documents": len(required),
        "missing_documents": missing_docs,
        "issues_found": all_issues
    }
    
    return output_docx_path, summary_report